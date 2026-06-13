"""DOCX parser — extracts text and metadata from Word documents."""

from pathlib import Path

from docx import Document as DocxDocument

from utils.logger import setup_logger

logger = setup_logger(__name__)


class DocxParser:
    """Handles DOCX document parsing and metadata extraction."""

    def parse(self, file_path: str) -> str:
        """Extract text from a DOCX file.

        Extracts text from paragraphs and tables.

        Args:
            file_path: Path to the DOCX file.

        Returns:
            Extracted text content.

        Raises:
            FileNotFoundError: If file doesn't exist.
            ValueError: If no text could be extracted.
        """
        path = Path(file_path)
        if not path.exists():
            raise FileNotFoundError(f"DOCX file not found: {file_path}")

        logger.info("Parsing DOCX: %s", path.name)

        try:
            doc = DocxDocument(file_path)
            parts: list[str] = []

            # Extract paragraphs
            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    parts.append(text)

            # Extract tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                    if row_text:
                        parts.append(row_text)

            text = "\n\n".join(parts)

            if not text.strip():
                raise ValueError(f"Could not extract text from DOCX: {path.name}")

            logger.info("Extracted %d characters from %s", len(text), path.name)
            return text

        except ValueError:
            raise
        except Exception as e:
            logger.error("DOCX parsing failed: %s", e)
            raise ValueError(f"Failed to parse DOCX file: {e}")

    def extract_metadata(self, file_path: str) -> dict:
        """Extract metadata from a DOCX file.

        Returns:
            Dict with keys: author, title, word_count, paragraph_count.
        """
        metadata = {
            "author": "",
            "title": "",
            "word_count": 0,
            "paragraph_count": 0,
        }
        try:
            doc = DocxDocument(file_path)
            core = doc.core_properties
            metadata["author"] = core.author or ""
            metadata["title"] = core.title or ""
            metadata["paragraph_count"] = len(doc.paragraphs)

            # Count words
            word_count = 0
            for para in doc.paragraphs:
                word_count += len(para.text.split())
            metadata["word_count"] = word_count

        except Exception as e:
            logger.warning("Could not extract DOCX metadata: %s", e)
        return metadata
